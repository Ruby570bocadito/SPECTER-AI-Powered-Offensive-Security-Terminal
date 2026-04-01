from __future__ import annotations

import json
import csv
from typing import Any, Dict, List, Optional

try:
    from specter.skills.base import BaseSkill  # type: ignore
except Exception:
    # Fallback in case the runtime path differs during tests
    class BaseSkill:  # type: ignore
        pass


class ReportSkill(BaseSkill):
    """Reporting toolkit for SPECTER.

    This skill provides executive and technical summaries of a session, a
    findings matrix, and export helpers for Markdown, JSON and CSV formats.
    """

    def __init__(self, session: Optional[Any] = None) -> None:
        self.session = session

    # Internal helpers
    def _extract_session_findings(self, session: Any) -> List[Dict[str, Any]]:
        # Look for common session attributes that may contain findings.
        candidate_attrs = ["findings", "report_findings", "vulnerabilities"]
        for attr in candidate_attrs:
            val = getattr(session, attr, None)
            if isinstance(val, list):
                # Heuristic: expect list[dict]
                if all(isinstance(item, dict) for item in val):
                    return val
        return []

    def generate_executive_summary(self, session: Any) -> str:
        findings = self._extract_session_findings(session)
        count = len(findings) if findings else 0
        if count == 0:
            return "Sesión sin hallazgos reportados."

        # Aggregate by severity if available
        severities: Dict[str, int] = {}
        for f in findings:
            sev = f.get("severity") or f.get("level") or "unknown"
            severities[sev] = severities.get(sev, 0) + 1

        parts = [f"{count} hallazgos identificados"]
        for sev, n in severities.items():
            parts.append(f"{sev}: {n}")
        return "; ".join(parts) + "."

    def generate_findings_matrix(self, session: Any) -> List[Dict[str, Any]]:
        findings = self._extract_session_findings(session)
        # Normalize into a simple matrix: list of dicts with common keys
        matrix: List[Dict[str, Any]] = []
        for f in findings:
            entry = {
                "id": f.get("id") or f.get("name") or "",
                "name": f.get("name") or f.get("title") or "",
                "host": f.get("host") or f.get("target") or "",
                "port": f.get("port"),
                "severity": f.get("severity") or f.get("level") or "",
                "description": f.get("description") or f.get("desc") or "",
                "remediation": f.get("remediation") or "",
            }
            matrix.append(entry)
        return matrix

    def export_markdown(self, data: Any) -> str:
        # Pretty export to Markdown. Supports list[dict] (table) or other types.
        if isinstance(data, list) and data and isinstance(data[0], dict):
            headers = list(data[0].keys())
            lines: List[str] = []
            # header
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("|" + "|".join(["---"] * len(headers)) + "|")
            for row in data:
                line = "| " + " | ".join(str(row.get(h, "")) for h in headers) + " |"
                lines.append(line)
            return "\n".join(lines)
        if isinstance(data, dict):
            return "```json\n" + json.dumps(data, indent=2, ensure_ascii=False) + "\n```"
        if isinstance(data, list):
            return "- " + "\n- ".join(str(x) for x in data)
        return str(data)

    def export_json(self, data: Any) -> str:
        return json.dumps(data, indent=2, ensure_ascii=False)

    def export_csv(self, data: Any) -> str:
        # Only supports list[dict]
        if not (isinstance(data, list) and data and isinstance(data[0], dict)):
            return ""
        headers = list(data[0].keys())
        lines: List[str] = [",".join(headers)]
        for row in data:
            lines.append(",".join(str(row.get(h, "")) for h in headers))
        return "\n".join(lines)

    # --- New CVSS/HTML/Timeline helpers ---
    def generate_cvss_report(self, findings: List[Dict[str, Any]]) -> str:
        """Generate a simple CVSS-style report from findings.

        Each finding may contain a nested 'cvss' dict with 'baseScore' and 'vector'.
        """
        if not findings:
            return "No CVSS findings."
        lines: List[str] = ["CVSS Report:"]
        for f in findings:
            cvss = f.get("cvss") or {}
            base = cvss.get("baseScore") if isinstance(cvss, dict) else None
            vector = cvss.get("vector") if isinstance(cvss, dict) else None
            if base is None:
                base = f.get("severity") or "N/A"
            lines.append(f"- {f.get('name') or f.get('type')}: CVSS Base={base}, Vector={vector or 'N/A'}")
        return "\n".join(lines)

    def export_html(self, data: Any) -> str:
        """Export data to a simple HTML representation."""
        if isinstance(data, list) and data and isinstance(data[0], dict):
            headers = list(data[0].keys())
            rows = []
            for row in data:
                rows.append("<tr>" + "".join(f"<td>{row.get(h, '')}</td>" for h in headers) + "</tr>")
            thead = "<tr>" + "".join(f"<th>{h}</th>" for h in headers) + "</tr>"
            return f"<table border=1><thead>{thead}</thead><tbody>{''.join(rows)}</tbody></table>"
        if isinstance(data, dict):
            return f"<pre>{json.dumps(data, indent=2, ensure_ascii=False)}</pre>"
        return f"<p>{str(data)}</p>"

    def generate_timeline(self, session: Any) -> str:
        """Generate a simple HTML timeline from findings in a session."""
        findings = self._extract_session_findings(session)
        if not findings:
            return "<p>No timeline available.</p>"
        # Basic HTML timeline: list items with timestamp if present
        items = []
        for f in findings:
            ts = f.get("timestamp") or f.get("time") or "N/A"
            desc = f.get("description") or f.get("name") or f.get("type")
            items.append(f"<li>{ts} - {desc}</li>")
        return f"<ul>{''.join(items)}</ul>"

    # BaseSkill abstract methods
    @property
    def name(self) -> str:
        return "report"

    @property
    def description(self) -> str:
        return "Generación de informes"

    @property
    def category(self) -> str:
        return "report"

    @property
    def risk_level(self):
        from specter.skills.base import RiskLevel
        return RiskLevel.PASIVE

    async def execute(self, action: str, params: dict) -> Any:
        from specter.skills.base import SkillResult
        return SkillResult(success=True, output=f"Report action: {action}")

    async def validate_params(self, action: str, params: dict) -> bool:
        return True

    def get_available_actions(self) -> list:
        return ["generate_executive_summary", "generate_technical_report", "generate_findings_matrix", "export_markdown", "export_json", "export_csv", "generate_cvss_report", "export_html", "generate_timeline"]

    def export_pdf(self, data: Any, template: str = "default") -> bytes:
        """Exporta hallazgos a PDF.
        
        Args:
            data: Hallazgos a exportar
            template: Plantilla a usar (default, executive, technical)
            
        Returns:
            Bytes del PDF generado
        """
        try:
            from weasyprint import HTML
            html_content = self._generate_pdf_html(data, template)
            return HTML(string=html_content).write_pdf()
        except ImportError:
            return b"PDF export requires weasyprint: pip install weasyprint"
        except Exception as e:
            return f"Error generating PDF: {e}".encode()
    
    def _generate_pdf_html(self, data: Any, template: str) -> str:
        """Genera HTML para PDF según plantilla"""
        if template == "executive":
            return self._generate_executive_html(data)
        elif template == "technical":
            return self._generate_technical_html(data)
        return self._generate_default_html(data)
    
    def _generate_default_html(self, data: Any) -> str:
        findings = data if isinstance(data, list) else [data]
        rows = ""
        for f in findings:
            rows += f"""
            <tr>
                <td>{f.get('id', '')}</td>
                <td>{f.get('title', f.get('name', ''))}</td>
                <td>{f.get('severity', '')}</td>
                <td>{f.get('cvss', '')}</td>
            </tr>
            """
        return f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ color: #00D4FF; }}
                table {{ width: 100%; border-collapse: collapse; }}
                th {{ background: #00FF88; color: white; padding: 10px; text-align: left; }}
                td {{ padding: 8px; border: 1px solid #ddd; }}
                .crit {{ color: #FF3366; font-weight: bold; }}
                .high {{ color: #FF6B35; }}
                .med {{ color: #FFD60A; }}
            </style>
        </head>
        <body>
            <h1>SPECTER - Report</h1>
            <table>
                <thead><tr><th>ID</th><th>Title</th><th>Severity</th><th>CVSS</th></tr></thead>
                <tbody>{rows}</tbody>
            </table>
        </body>
        </html>
        """
    
    def _generate_executive_html(self, data: Any) -> str:
        return f"""
        <html><head><style>
            body {{ font-family: Arial; margin: 40px; }}
            .header {{ background: #080C14; color: #00FF88; padding: 20px; }}
            h1 {{ color: #00FF88; }}
            .summary {{ background: #f5f5f5; padding: 15px; margin: 20px 0; }}
        </style></head><body>
            <div class="header"><h1>Executive Summary</h1></div>
            <div class="summary">
                <p>Generated by SPECTER AI-Powered Security Terminal</p>
            </div>
            {self.export_html(data)}
        </body></html>
        """
    
    def _generate_technical_html(self, data: Any) -> str:
        return f"""
        <html><head><style>
            body {{ font-family: 'Courier New', monospace; margin: 40px; }}
            h1 {{ color: #00D4FF; border-bottom: 2px solid #00FF88; }}
            pre {{ background: #f0f0f0; padding: 15px; }}
        </style></head><body>
            <h1>Technical Report</h1>
            <pre>{self.export_json(data)}</pre>
        </body></html>
        """

    def export_docx(self, findings: list, title: str = "SPECTER Report") -> bytes:
        """Exporta hallazgos a DOCX.
        
        Args:
            findings: Lista de hallazgos
            title: Título del documento
            
        Returns:
            Bytes del documento DOCX
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph("Generated by SPECTER - AI-Powered Offensive Security Terminal")
            doc.add_paragraph()
            
            for i, f in enumerate(findings, 1):
                doc.add_heading(f"{i}. {f.get('title', f.get('name', 'Unknown'))}", level=2)
                
                severity = f.get('severity', 'INFO')
                para = doc.add_paragraph(f"Severity: {severity}")
                if severity == "CRIT":
                    para.runs[0].font.color.rgb = RGBColor(255, 51, 102)
                elif severity == "HIGH":
                    para.runs[0].font.color.rgb = RGBColor(255, 107, 53)
                
                if f.get('cvss'):
                    doc.add_paragraph(f"CVSS: {f['cvss']}")
                
                if f.get('description'):
                    doc.add_paragraph(f"Description: {f['description']}")
                
                if f.get('remediation'):
                    doc.add_paragraph(f"Remediation: {f['remediation']}")
                
                doc.add_paragraph()
            
            import io
            f = io.BytesIO()
            doc.save(f)
            return f.getvalue()
        except ImportError:
            return b"DOCX export requires python-docx: pip install python-docx"
        except Exception as e:
            return f"Error generating DOCX: {e}".encode()
